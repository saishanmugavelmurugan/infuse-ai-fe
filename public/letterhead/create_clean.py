#!/usr/bin/env python3
"""
Create clean editable letterhead - NO borders/outlines
"""

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/app/frontend/public/letterhead"

AMBER = "#F9A825"
ORANGE = "#F57C00" 
RED = "#E53935"

def hex_to_rgb_tuple(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def create_clean_header():
    """Create header without any border"""
    original_path = os.path.join(OUTPUT_DIR, "original-letterhead.png")
    img = Image.open(original_path).convert('RGB')  # RGB not RGBA to avoid transparency issues
    width, height = img.size
    
    # Crop header portion (logo + text + line)
    header_height = int(height * 0.155)
    header_img = img.crop((0, 0, width, header_height))
    
    header_path = os.path.join(OUTPUT_DIR, "clean-header.png")
    header_img.save(header_path, "PNG")
    return header_path, width

def create_clean_footer(width):
    """Create footer with straight lines and address below"""
    footer_height = 85
    footer_img = Image.new('RGB', (width, footer_height), (255, 255, 255))
    draw = ImageDraw.Draw(footer_img)
    
    # Three straight lines (no gaps)
    line_height = 16
    y = 0
    
    draw.rectangle([0, y, width, y + line_height], fill=hex_to_rgb_tuple(AMBER))
    y += line_height
    draw.rectangle([0, y, width, y + line_height], fill=hex_to_rgb_tuple(ORANGE))
    y += line_height
    draw.rectangle([0, y, width, y + line_height], fill=hex_to_rgb_tuple(RED))
    
    # Address below in fine print
    y += line_height + 10
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/freefont/FreeSans.ttf", 13)
        address = "1 G-Floor Tower P1 Sec 65, Emaar Emerald Estate, Badshahpur, Gurgaon- 122101, Haryana"
        bbox = draw.textbbox((0, 0), address, font=font)
        text_w = bbox[2] - bbox[0]
        draw.text((width - text_w - 15, y), address, fill=(100, 100, 100), font=font)
    except:
        pass
    
    footer_path = os.path.join(OUTPUT_DIR, "clean-footer.png")
    footer_img.save(footer_path, "PNG")
    return footer_path

def remove_picture_border(run):
    """Remove border from inline picture"""
    # Get the inline element
    inline = run._r.xpath('.//a:blip/..')[0] if run._r.xpath('.//a:blip/..') else None
    if inline is not None:
        # Try to remove any border/outline properties
        pass

def create_letterhead():
    """Create clean letterhead document"""
    
    header_path, img_width = create_clean_header()
    footer_path = create_clean_footer(img_width)
    
    doc = Document()
    
    # Page setup - A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.3)
    
    # HEADER
    header = section.header
    # Clear default paragraph
    header.paragraphs[0].clear()
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    
    run = hp.add_run()
    picture = run.add_picture(header_path, width=Cm(17))
    
    # FOOTER
    footer = section.footer
    footer.paragraphs[0].clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    
    footer_run = fp.add_run()
    footer_run.add_picture(footer_path, width=Cm(17))
    
    # Save .docx
    docx_path = os.path.join(OUTPUT_DIR, "Infuse-Letterhead-Clean.docx")
    doc.save(docx_path)
    print(f"Created: {docx_path}")
    
    return docx_path

def create_continuation():
    """Create continuation sheet"""
    
    original_path = os.path.join(OUTPUT_DIR, "original-letterhead.png")
    img = Image.open(original_path)
    img_width = img.size[0]
    
    footer_path = os.path.join(OUTPUT_DIR, "clean-footer.png")
    
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.footer_distance = Cm(0.3)
    
    # Footer only
    footer = section.footer
    footer.paragraphs[0].clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_run = fp.add_run()
    footer_run.add_picture(footer_path, width=Cm(17))
    
    cont_path = os.path.join(OUTPUT_DIR, "Infuse-Continuation-Clean.docx")
    doc.save(cont_path)
    print(f"Created: {cont_path}")
    
    return cont_path

if __name__ == "__main__":
    print("Creating clean letterhead (no borders)...")
    create_letterhead()
    create_continuation()
    print("\nDone!")
