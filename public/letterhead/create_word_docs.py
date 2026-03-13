#!/usr/bin/env python3
"""
Create Word documents for letterhead - NO watermark
"""

from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/app/frontend/public/letterhead"

# Brand colors
AMBER = "#F9A825"
ORANGE = "#F57C00" 
RED = "#E53935"
WHITE = "#FFFFFF"

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def create_png_without_watermark():
    """Create PNG version without watermark"""
    original_path = os.path.join(OUTPUT_DIR, "original-letterhead.png")
    img = Image.open(original_path).convert('RGBA')
    width, height = img.size
    
    draw = ImageDraw.Draw(img)
    
    # Cover the wavy area with white
    wavy_start_y = int(height * 0.87)
    draw.rectangle([0, wavy_start_y, width, height], fill=(255, 255, 255, 255))
    
    # Redraw address
    try:
        from PIL import ImageFont
        font_address = ImageFont.truetype("/usr/share/fonts/truetype/freefont/FreeSansBold.ttf", 28)
        address_lines = [
            "1 G-Floor Tower P1 Sec 65,",
            "Emaar Emerald Estate, Badshahpur,",
            "Gurgaon- 122101, Haryana"
        ]
        
        address_y = wavy_start_y + 20
        address_margin = 60
        
        for i, line in enumerate(address_lines):
            bbox = draw.textbbox((0, 0), line, font=font_address)
            line_w = bbox[2] - bbox[0]
            draw.text((width - address_margin - line_w, address_y + i * 38), line, 
                      fill=(51, 51, 51, 255), font=font_address)
    except:
        pass
    
    # Draw THREE STRAIGHT LINES at bottom
    line_height = 22
    bottom_margin = 10
    
    line_y_red = height - bottom_margin - line_height
    line_y_orange = line_y_red - line_height
    line_y_amber = line_y_orange - line_height
    
    draw.rectangle([0, line_y_amber, width, line_y_amber + line_height], fill=hex_to_rgb(AMBER) + (255,))
    draw.rectangle([0, line_y_orange, width, line_y_orange + line_height], fill=hex_to_rgb(ORANGE) + (255,))
    draw.rectangle([0, line_y_red, width, line_y_red + line_height], fill=hex_to_rgb(RED) + (255,))
    
    # Save without watermark
    output_path = os.path.join(OUTPUT_DIR, "letterhead-no-watermark.png")
    img.save(output_path, "PNG")
    print(f"Created PNG: {output_path}")
    return output_path

def add_colored_line(paragraph, color_hex):
    """Add a colored horizontal line"""
    run = paragraph.add_run()
    # Create a shape for the line
    
def create_word_letterhead():
    """Create Word document letterhead"""
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Add the header image (use the PNG we created)
    header_img_path = os.path.join(OUTPUT_DIR, "letterhead-no-watermark.png")
    
    # Create a header section
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add header image
    run = header_para.add_run()
    run.add_picture(header_img_path, width=Inches(6.5))
    
    # Add content area (empty paragraphs for writing space)
    for _ in range(25):
        doc.add_paragraph()
    
    # Save
    output_path = os.path.join(OUTPUT_DIR, "Infuse-AI-Letterhead.docx")
    doc.save(output_path)
    print(f"Created Word: {output_path}")
    return output_path

def create_word_letterhead_v2():
    """Create Word document with embedded elements"""
    doc = Document()
    
    # Set page margins (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # === HEADER SECTION ===
    # Add logo image at top
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Use the original letterhead image for header portion only
    # First create a cropped header image
    original_path = os.path.join(OUTPUT_DIR, "original-letterhead.png")
    img = Image.open(original_path)
    width, height = img.size
    
    # Crop just the header portion (top ~15%)
    header_height = int(height * 0.16)
    header_img = img.crop((0, 0, width, header_height))
    header_img_path = os.path.join(OUTPUT_DIR, "header-only.png")
    header_img.save(header_img_path)
    
    run = logo_para.add_run()
    run.add_picture(header_img_path, width=Inches(6))
    
    # Add horizontal line after header
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # === CONTENT AREA ===
    # Add empty paragraphs for content
    for _ in range(20):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(18)
    
    # === FOOTER with address and lines ===
    # Address (right-aligned)
    address_para = doc.add_paragraph()
    address_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    address_run = address_para.add_run("1 G-Floor Tower P1 Sec 65,\nEmaar Emerald Estate, Badshahpur,\nGurgaon- 122101, Haryana")
    address_run.font.size = Pt(10)
    address_run.font.bold = True
    address_run.font.color.rgb = RGBColor(51, 51, 51)
    
    # Create footer lines image
    lines_img = Image.new('RGB', (2000, 70), (255, 255, 255))
    lines_draw = ImageDraw.Draw(lines_img)
    lines_draw.rectangle([0, 0, 2000, 22], fill=hex_to_rgb(AMBER))
    lines_draw.rectangle([0, 23, 2000, 45], fill=hex_to_rgb(ORANGE))
    lines_draw.rectangle([0, 46, 2000, 68], fill=hex_to_rgb(RED))
    lines_path = os.path.join(OUTPUT_DIR, "footer-lines.png")
    lines_img.save(lines_path)
    
    # Add footer lines
    lines_para = doc.add_paragraph()
    lines_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines_run = lines_para.add_run()
    lines_run.add_picture(lines_path, width=Inches(7.5))
    
    # Save
    output_path = os.path.join(OUTPUT_DIR, "Infuse-AI-Letterhead.docx")
    doc.save(output_path)
    print(f"Created Word: {output_path}")
    return output_path

def create_continuation_word():
    """Create continuation sheet Word document"""
    doc = Document()
    
    # Set page margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # Add thin line at top
    top_line = Image.new('RGB', (2000, 3), hex_to_rgb("#333333"))
    top_line_path = os.path.join(OUTPUT_DIR, "top-line.png")
    top_line.save(top_line_path)
    
    top_para = doc.add_paragraph()
    top_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    top_run = top_para.add_run()
    top_run.add_picture(top_line_path, width=Inches(6.5))
    
    # Content area
    for _ in range(22):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(18)
    
    # Footer lines
    lines_path = os.path.join(OUTPUT_DIR, "footer-lines.png")
    lines_para = doc.add_paragraph()
    lines_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines_run = lines_para.add_run()
    lines_run.add_picture(lines_path, width=Inches(7.5))
    
    # Save
    output_path = os.path.join(OUTPUT_DIR, "Infuse-AI-Continuation.docx")
    doc.save(output_path)
    print(f"Created Word: {output_path}")
    return output_path

if __name__ == "__main__":
    print("Creating letterhead files...")
    create_png_without_watermark()
    create_word_letterhead_v2()
    create_continuation_word()
    print("\nAll files created!")
