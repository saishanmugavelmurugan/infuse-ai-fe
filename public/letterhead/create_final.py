#!/usr/bin/env python3
"""
Keep EXACT original design, only change:
1. Wavy lines -> Straight lines
2. Address below lines in smaller print
"""

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = "/app/frontend/public/letterhead"

# Brand colors from original
AMBER = "#F9A825"
ORANGE = "#F57C00" 
RED = "#E53935"

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def create_final_letterhead():
    """Create the exact letterhead as requested"""
    original_path = os.path.join(OUTPUT_DIR, "original-letterhead.png")
    img = Image.open(original_path).convert('RGBA')
    width, height = img.size
    print(f"Image size: {width}x{height}")
    
    draw = ImageDraw.Draw(img)
    
    # Find where wavy design starts and cover it with white
    wavy_start_y = int(height * 0.855)
    draw.rectangle([0, wavy_start_y, width, height], fill=(255, 255, 255, 255))
    
    # Draw THREE STRAIGHT LINES
    line_height = 20
    line_gap = 0
    lines_start_y = wavy_start_y + 30
    
    # Amber line (top)
    draw.rectangle([0, lines_start_y, width, lines_start_y + line_height], 
                   fill=hex_to_rgb(AMBER) + (255,))
    # Orange line (middle)
    draw.rectangle([0, lines_start_y + line_height + line_gap, width, lines_start_y + 2*line_height + line_gap], 
                   fill=hex_to_rgb(ORANGE) + (255,))
    # Red line (bottom)
    draw.rectangle([0, lines_start_y + 2*line_height + 2*line_gap, width, lines_start_y + 3*line_height + 2*line_gap], 
                   fill=hex_to_rgb(RED) + (255,))
    
    # Address BELOW lines in finer print (right-aligned)
    address_y = lines_start_y + 3*line_height + 15
    
    try:
        font_address = ImageFont.truetype("/usr/share/fonts/truetype/freefont/FreeSans.ttf", 18)
        address_text = "1 G-Floor Tower P1 Sec 65, Emaar Emerald Estate, Badshahpur, Gurgaon- 122101, Haryana"
        
        bbox = draw.textbbox((0, 0), address_text, font=font_address)
        text_w = bbox[2] - bbox[0]
        draw.text((width - text_w - 30, address_y), address_text, 
                  fill=(80, 80, 80, 255), font=font_address)
    except Exception as e:
        print(f"Font error: {e}")
    
    # Save PNG
    png_path = os.path.join(OUTPUT_DIR, "final-letterhead.png")
    img.save(png_path, "PNG")
    print(f"Created: {png_path}")
    
    # Create continuation sheet (just lines + address, no header)
    cont_img = Image.new('RGBA', (width, height), (255, 255, 255, 255))
    cont_draw = ImageDraw.Draw(cont_img)
    
    # Same lines at same position
    cont_draw.rectangle([0, lines_start_y, width, lines_start_y + line_height], 
                       fill=hex_to_rgb(AMBER) + (255,))
    cont_draw.rectangle([0, lines_start_y + line_height, width, lines_start_y + 2*line_height], 
                       fill=hex_to_rgb(ORANGE) + (255,))
    cont_draw.rectangle([0, lines_start_y + 2*line_height, width, lines_start_y + 3*line_height], 
                       fill=hex_to_rgb(RED) + (255,))
    
    # Same address
    try:
        cont_draw.text((width - text_w - 30, address_y), address_text, 
                      fill=(80, 80, 80, 255), font=font_address)
    except:
        pass
    
    cont_path = os.path.join(OUTPUT_DIR, "final-continuation.png")
    cont_img.save(cont_path, "PNG")
    print(f"Created: {cont_path}")
    
    return png_path, cont_path

def create_word_docs():
    """Create Word documents using the PNG images"""
    
    # Main letterhead
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)
    
    # Add the letterhead image as background
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(os.path.join(OUTPUT_DIR, "final-letterhead.png"), width=Inches(8.27))
    
    docx_path = os.path.join(OUTPUT_DIR, "Letterhead.docx")
    doc.save(docx_path)
    print(f"Created: {docx_path}")
    
    # Continuation sheet
    doc2 = Document()
    section2 = doc2.sections[0]
    section2.page_width = Cm(21)
    section2.page_height = Cm(29.7)
    section2.top_margin = Cm(0)
    section2.bottom_margin = Cm(0)
    section2.left_margin = Cm(0)
    section2.right_margin = Cm(0)
    
    para2 = doc2.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = para2.add_run()
    run2.add_picture(os.path.join(OUTPUT_DIR, "final-continuation.png"), width=Inches(8.27))
    
    cont_docx_path = os.path.join(OUTPUT_DIR, "Continuation.docx")
    doc2.save(cont_docx_path)
    print(f"Created: {cont_docx_path}")

if __name__ == "__main__":
    print("Creating letterheads...")
    create_final_letterhead()
    create_word_docs()
    print("\nDone!")
