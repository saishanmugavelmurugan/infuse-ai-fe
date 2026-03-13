#!/usr/bin/env python3
"""
Create EDITABLE Word letterhead with:
- Header section: logo + company details
- Footer section: straight lines + address
- Body: empty and editable for writing
"""

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/app/frontend/public/letterhead"

# Brand colors
AMBER = "#F9A825"
ORANGE = "#F57C00" 
RED = "#E53935"

def hex_to_rgb_tuple(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def create_header_image():
    """Create just the header portion from original"""
    original_path = os.path.join(OUTPUT_DIR, "original-letterhead.png")
    img = Image.open(original_path).convert('RGBA')
    width, height = img.size
    
    # Crop header (top portion with logo, name, details, line)
    header_height = int(height * 0.155)
    header_img = img.crop((0, 0, width, header_height))
    
    header_path = os.path.join(OUTPUT_DIR, "header-image.png")
    header_img.save(header_path, "PNG")
    print(f"Created header: {header_path}")
    return header_path, width

def create_footer_image(width):
    """Create footer with straight lines"""
    footer_height = 90
    footer_img = Image.new('RGBA', (width, footer_height), (255, 255, 255, 255))
    draw = ImageDraw.Draw(footer_img)
    
    # Three straight lines
    line_height = 18
    y = 0
    
    draw.rectangle([0, y, width, y + line_height], fill=hex_to_rgb_tuple(AMBER) + (255,))
    y += line_height
    draw.rectangle([0, y, width, y + line_height], fill=hex_to_rgb_tuple(ORANGE) + (255,))
    y += line_height
    draw.rectangle([0, y, width, y + line_height], fill=hex_to_rgb_tuple(RED) + (255,))
    
    # Address below in fine print
    y += line_height + 8
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/freefont/FreeSans.ttf", 14)
        address = "1 G-Floor Tower P1 Sec 65, Emaar Emerald Estate, Badshahpur, Gurgaon- 122101, Haryana"
        bbox = draw.textbbox((0, 0), address, font=font)
        text_w = bbox[2] - bbox[0]
        draw.text((width - text_w - 20, y), address, fill=(100, 100, 100, 255), font=font)
    except:
        pass
    
    footer_path = os.path.join(OUTPUT_DIR, "footer-image.png")
    footer_img.save(footer_path, "PNG")
    print(f"Created footer: {footer_path}")
    return footer_path

def create_editable_letterhead():
    """Create editable Word document with header/footer"""
    
    # First create header and footer images
    header_path, img_width = create_header_image()
    footer_path = create_footer_image(img_width)
    
    # Create Word document
    doc = Document()
    
    # Set page size to A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(4.5)  # Space for header
    section.bottom_margin = Cm(3)  # Space for footer
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # === HEADER ===
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_after = Pt(0)
    header_para.paragraph_format.space_before = Pt(0)
    
    run = header_para.add_run()
    run.add_picture(header_path, width=Cm(18))
    
    # === FOOTER ===
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_after = Pt(0)
    footer_para.paragraph_format.space_before = Pt(0)
    
    footer_run = footer_para.add_run()
    footer_run.add_picture(footer_path, width=Cm(18))
    
    # === BODY (editable area) ===
    # Add some empty paragraphs so user can type
    # The body is automatically editable
    
    # Save as .docx
    docx_path = os.path.join(OUTPUT_DIR, "Infuse-Letterhead-Editable.docx")
    doc.save(docx_path)
    print(f"Created: {docx_path}")
    
    # Also save as .dotx template
    dotx_path = os.path.join(OUTPUT_DIR, "Infuse-Letterhead-Template.dotx")
    doc.save(dotx_path)
    print(f"Created: {dotx_path}")
    
    return docx_path, dotx_path

def create_continuation_sheet():
    """Create continuation sheet (no header, just footer)"""
    
    # Get image width from original
    original_path = os.path.join(OUTPUT_DIR, "original-letterhead.png")
    img = Image.open(original_path)
    img_width = img.size[0]
    
    footer_path = os.path.join(OUTPUT_DIR, "footer-image.png")
    
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Footer only
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run()
    footer_run.add_picture(footer_path, width=Cm(18))
    
    # Save
    cont_path = os.path.join(OUTPUT_DIR, "Infuse-Continuation-Editable.docx")
    doc.save(cont_path)
    print(f"Created: {cont_path}")
    
    return cont_path

if __name__ == "__main__":
    print("Creating editable letterhead...")
    create_editable_letterhead()
    create_continuation_sheet()
    print("\nDone! You can now type in these documents.")
